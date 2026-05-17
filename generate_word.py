import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        
def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

doc = docx.Document()

# Title
title = doc.add_heading('Jawaban Assessment - UMKM Insight (RPL-SMS_4)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Soal 1
add_heading(doc, 'SOAL 1: Nama & Deskripsi Aplikasi')
p = doc.add_paragraph()
p.add_run('Nama Aplikasi: ').bold = True
p.add_run('UMKM Insight\n\n')
p.add_run('Deskripsi:\n').bold = True
p.add_run('UMKM Insight adalah aplikasi SaaS (Software as a Service) berbasis web yang berfungsi sebagai platform ')
p.add_run('Read-Only Analytics').bold = True
p.add_run(' untuk pelaku UMKM. Aplikasi ini menarik data transaksi keuangan dari sistem perbankan SmartBank melalui API Gateway, lalu mengolahnya menjadi dashboard analitik berupa ringkasan penjualan, tren bulanan, distribusi transaksi, breakdown biaya operasional, serta insight bisnis otomatis. UMKM Insight ')
p.add_run('tidak memproses uang').bold = True
p.add_run(' — seluruh aliran dana tetap melalui SmartBank. Aplikasi menerapkan model monetisasi berlangganan bertingkat (Free/Basic/Pro/Enterprise) di mana pembayaran langganan juga diproses oleh SmartBank, dan fitur yang tampil di dashboard dibatasi sesuai paket langganan pengguna.\n\n')
p.add_run('Tech Stack: ').bold = True
p.add_run('Frontend (Vanilla JS + Vite + Chart.js + Axios), Backend (Python Flask REST API), Database (MySQL via XAMPP), Data Transaksi (JSON file simulasi SmartBank).')

# Soal 2
add_heading(doc, 'SOAL 2 (Bobot 50): Analisis Proses Transaksi End-to-End')
add_heading(doc, 'Konteks Peran', level=2)
doc.add_paragraph('UMKM Insight adalah aplikasi analytics read-only dalam ekosistem UMKM digital. Aplikasi ini berinteraksi dengan SmartBank (untuk pembayaran langganan & pelaporan performa) dan menerima data dari API Gateway yang menghubungkan ke Marketplace, POS, SupplierHub, dan LogistikKita.')

add_heading(doc, '1) Input Utama yang Diterima Aplikasi', level=2)
table1 = doc.add_table(rows=1, cols=4)
table1.style = 'Table Grid'
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Input'
hdr_cells[1].text = 'Sumber'
hdr_cells[2].text = 'Endpoint'
hdr_cells[3].text = 'Keterangan'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True

row_data1 = [
    ('Email + Password', 'User (Frontend)', 'POST /api/auth/login', 'Untuk autentikasi'),
    ('user_id (UMKM ID)', 'Query parameter', 'GET /api/umkm_insight/dashboard', 'Identifier UMKM'),
    ('Data transaksi mentah', 'File dummy_data.json', 'Internal get_transactions()', 'Berisi amount, fee, dll'),
    ('Pilihan paket + durasi', 'User (Frontend)', 'GET /api/smartbank/pay', 'Untuk upgrade langganan'),
    ('Laporan performa', 'User (Frontend)', 'POST /api/smartbank/report-decline', 'JSON body laporan')
]
for r in row_data1:
    row_cells = table1.add_row().cells
    for i in range(4):
        row_cells[i].text = r[i]

add_heading(doc, '2) API yang Perlu Dipanggil ke Sistem Lain', level=2)
table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'API Call'
hdr2[1].text = 'Target'
hdr2[2].text = 'Method'
hdr2[3].text = 'Tujuan'
for cell in hdr2:
    cell.paragraphs[0].runs[0].bold = True

row_data2 = [
    ('/api/smartbank/pay', 'SmartBank (Simulasi)', 'GET', 'Memproses pembayaran langganan'),
    ('/api/smartbank/report-decline', 'SmartBank (Simulasi)', 'POST', 'Mengirim laporan performa'),
    ('get_transactions(user_id)', 'SmartBank Data (JSON)', 'Internal', 'Mengambil data transaksi read-only'),
    ('get_smartbank_user_profile', 'SmartBank Data (JSON)', 'Internal', 'Mengambil profil UMKM'),
    ('get_fee_structure()', 'SmartBank Data (JSON)', 'Internal', 'Mengambil struktur fee')
]
for r in row_data2:
    row_cells = table2.add_row().cells
    for i in range(4):
        row_cells[i].text = r[i]

doc.add_paragraph('Dalam konteks nyata, data transaksi berasal dari: Marketplace (fee 2%), POS (fee 1%), SupplierHub (fee 3%), LogistikKita (fee 5%). Semua dijembatani SmartBank.')

add_heading(doc, '3) Data yang Dikirim dan Diterima', level=2)
doc.add_paragraph('A. Alur Dashboard (Aktivitas Ekonomi Utama):').bold = True
p = doc.add_paragraph('Backend memproses:\n1. Ambil profil SmartBank user\n2. Ambil subscription dari MySQL\n3. Ambil transaksi dari JSON\n4. Jalankan analisis_lengkap()\n5. Jalankan enrich_analytics()\n6. Jalankan limit_dashboard_by_package()\n\nResponse JSON ke Frontend: user data, analisis (ringkasan, tren, performa, dll), transaksi terbaru.')

doc.add_paragraph('B. Alur Pembayaran Langganan:').bold = True
p = doc.add_paragraph('Data dikirim: package_name, duration, amount.\nSmartBank merespon: status success, ref, amount, timestamp.\nBackend kemudian meng-update tabel subscriptions dan mencatat ke payment_logs.')

add_heading(doc, '4) Mekanisme Validasi JWT/Token', level=2)
p = doc.add_paragraph('Frontend menyimpan JWT di LocalStorage setelah login. Setiap request API mengirim header ')
p.add_run('Authorization: Bearer <token>').bold = True
p.add_run('. Middleware @require_auth mendecode token menggunakan JWT_SECRET. Jika expired/invalid, akan mengembalikan error 401. Jika valid, payload (id, email, role, umkm_id) disuntikkan ke request.')

add_heading(doc, '5) Risiko Inkonsistensi Data', level=2)
add_bullet(doc, 'Race condition pada upgrade langganan (Double charge).')
add_bullet(doc, 'Data transaksi stale (Karena membaca dari file JSON lokal).')
add_bullet(doc, 'Subscription expiry check timing (SmartBank delay menyebabkan status tidak sinkron).')
add_bullet(doc, 'Inconsistency antara payment_logs dan subscriptions jika MySQL commit gagal.')
add_bullet(doc, 'UMKM ID collision jika 2 user register bersamaan.')

add_heading(doc, '6) Dampak Jika Salah Satu Aplikasi Lain Kegagalan', level=2)
add_bullet(doc, 'SmartBank down: Pembayaran dan pelaporan gagal. Namun dashboard tetap bisa diakses karena data transaksi dibaca dari JSON lokal.')
add_bullet(doc, 'Marketplace/SupplierHub/LogistikKita/POS down: Data transaksi baru tidak masuk. Dashboard menampilkan data historis lama (stale data).')

add_heading(doc, '7) Strategi Agar Sistem Tetap Robust', level=2)
add_bullet(doc, 'Idempotency pada pembayaran menggunakan smartbank_ref unik.')
add_bullet(doc, 'Database transaction dengan rollback otomatis.')
add_bullet(doc, 'Graceful degradation (Sistem tidak crash saat sumber data lain down).')
add_bullet(doc, 'Subscription expiry auto-check pada setiap request.')

# Soal 3
doc.add_page_break()
add_heading(doc, 'SOAL 3 (Bobot 50): Analisis Respons terhadap Lonjakan & Gangguan')
add_heading(doc, 'Analisis Respons UMKM Insight', level=2)

add_bullet(doc, '1) Transaksi Ekonomi Tetap Konsisten: Aplikasi menerapkan prinsip Read-Only. Tidak ada operasi INSERT/UPDATE/DELETE terhadap data transaksi keuangan.')
add_bullet(doc, '2) Tidak Terjadi Double Transaction: Setiap pembayaran menggunakan smartbank_ref unik dan constraint database memastikan tidak ada subscription ganda.')
add_bullet(doc, '3) Tidak Terjadi Pengurangan Stok Palsu: UMKM Insight tidak mengelola stok, hanya menampilkan data transaksi yang sudah "success" dari SupplierHub.')
add_bullet(doc, '4) Sistem Tetap Scalable: Menggunakan connection pooling MySQL dan API yang stateless.')
add_bullet(doc, '5) User Tetap Mendapatkan Feedback Jelas: Menampilkan HTTP Status dan pesan error yang spesifik (misal 401 Unauthorized, 400 Bad Request).')
add_bullet(doc, '6) Ekosistem Tidak Mengalami Cascade Failure: Aplikasi terisolasi. Jika layanan lain delay/down, dashboard UMKM Insight tetap hidup menampilkan data historis lokal.')

add_heading(doc, 'Komponen Kritis dan Prioritas', level=2)
doc.add_paragraph('1) Komponen Paling Kritis:').bold = True
add_bullet(doc, 'Middleware Auth (JWT) - Mengamankan seluruh endpoint.')
add_bullet(doc, 'Database Connection Pool - Kunci kelancaran read/write ke MySQL.')

doc.add_paragraph('2) Endpoint/API Prioritas:').bold = True
add_bullet(doc, 'P0 (Critical): /api/auth/login dan /api/umkm_insight/dashboard')
add_bullet(doc, 'P1 (High): /api/smartbank/pay')

doc.add_paragraph('3) Log yang Wajib Dicatat:').bold = True
add_bullet(doc, 'Request/Response log (API Gateway logger).')
add_bullet(doc, 'Activity Logs (Login, Upgrade, Expired, Report).')
add_bullet(doc, 'Payment Logs (Ref pembayaran dan status dari gateway).')

doc.add_paragraph('4) Prinsip Clean Architecture / SOLID:').bold = True
add_bullet(doc, 'Single Responsibility Principle (SRP): Modul terpisah sesuai fungsi (auth_routes.py, analysis.py, dll) sehingga saat lonjakan, scaling bisa difokuskan per fungsi.')
add_bullet(doc, 'Dependency Inversion Principle (DIP): Modul data_store.py mengabstraksi sumber data, sehingga perubahan dari JSON file ke REST API SmartBank kelak tidak merusak logic dashboard.')
add_bullet(doc, 'Separation of Concerns: Memisahkan UI, Auth, Business Logic, dan Data Access Layer mencegah cascade failure.')

doc.save('Jawaban_Assessment_UMKM_Insight.docx')
print("Document generated successfully.")
